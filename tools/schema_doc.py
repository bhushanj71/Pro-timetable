"""Generate the database and logging reference as a Word document.

Read from SQLAlchemy's live metadata rather than typed out by hand, because a
schema document that is maintained separately from the schema is wrong within a
month and nobody notices until they trust it. Re-run this after any model
change:

    python tools/schema_doc.py [output.docx]

The prose is curated -- purposes, the reasons behind the odd decisions, the
logging behaviour -- because introspection can list a column but cannot say why
it is there. The structure around it is generated, so the two cannot disagree
about what exists.
"""
from __future__ import annotations

import os
import sys
from datetime import date
from pathlib import Path

# Run from anywhere: tools/ is not the package root, and this script is meant to
# be runnable as `python tools/schema_doc.py` from the repository root or not.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault("DATABASE_URL", "sqlite:///./profschedule.db")

from docx import Document                                    # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH                # noqa: E402
from docx.shared import Inches, Pt                           # noqa: E402

import app.models  # noqa: F401,E402  -- registers every table on the metadata
from app.database import _ADDITIVE_COLUMNS, Base             # noqa: E402
from app.services.doc_export import _shade, _table_borders, _write_lines  # noqa: E402

HEAD_SHADE = "D9D9D9"
GROUP_SHADE = "EFEFEF"

# --------------------------------------------------------------------------
# Curated prose. Introspection knows a column exists; only this knows why.
# --------------------------------------------------------------------------
GROUPS: list[tuple[str, str, list[str]]] = [
    (
        "Identity and profile",
        "One account per professor. There is no separate 'work account': the same "
        "row carries the personal schedule, the organisational placement and the "
        "planning preferences, and `active_profile` only decides which view opens "
        "first.",
        ["users"],
    ),
    (
        "Organisation",
        "Which college somebody belongs to, and which part of it. Administrative "
        "posts such as Principal and Registrar are rows in `departments` "
        "distinguished by `kind`, because they answer the same question a "
        "department does and every membership and directory lookup already runs "
        "on a department id.",
        ["colleges", "departments"],
    ),
    (
        "Personal schedule",
        "The professor's own timetable, private to them. Every query in this group "
        "is scoped by `user_id`; nothing here is ever shared with a community.",
        ["events", "tasks", "reminders"],
    ),
    (
        "Work: communities and tasks",
        "Shared work. A task belongs to a community and is created by one person, "
        "but responsibility lives on `task_assignments`: one row per person per "
        "task, each with its own status and progress. Three people can share a "
        "task and be in three different states.",
        ["communities", "community_members", "community_invitations",
         "work_tasks", "task_assignments"],
    ),
    (
        "Work records and audit trail",
        "What happened to a task, and the evidence for it. These are the tables "
        "that answer 'how did this progress over time' and 'what was submitted'. "
        "See the Logs section for how they combine into one timeline.",
        ["task_progress_updates", "task_activities", "task_comments",
         "task_attachments"],
    ),
    (
        "Notifications and delivery",
        "The notification bell and the push endpoints behind it. "
        "`work_notifications` carries both work events and personal receipts "
        "despite its name -- it was already the bell's backing store, with read "
        "state and dismissal, and a second table would have meant a second feed "
        "to merge.",
        ["work_notifications", "push_subscriptions"],
    ),
    (
        "AI",
        "Transcript of the assistant's conversations, kept per user.",
        ["ai_conversations"],
    ),
]

PURPOSE: dict[str, str] = {
    "users": "The account. Also holds working hours, meal windows, commute and "
             "study preferences, which the day planner reads as preferences "
             "rather than rules.",
    "colleges": "An institution. Seeded with one default; administrators may add more.",
    "departments": "A teaching department or an administrative post, told apart "
                   "by `kind` ('academic' or 'office').",
    "events": "One occurrence on the timetable. A recurring lecture is many rows "
              "sharing a `recurrence_group_id`, not one row with a rule, so a "
              "single week can be cancelled without touching the rest.",
    "tasks": "A personal to-do, unrelated to communities. Distinct from "
             "`work_tasks`, which is shared work.",
    "reminders": "A promise to interrupt someone at a time: push, email, or both. "
                 "Deliberately separate from notifications, which are receipts.",
    "communities": "A team. Owned by its creator, who may delete it.",
    "community_members": "Membership and role: owner, admin or member. This row is "
                         "what every Work permission check reads.",
    "community_invitations": "A pending request to join. Membership is created only "
                             "on acceptance.",
    "work_tasks": "A shared task. Its `status` is derived from its assignments "
                  "rather than set directly.",
    "task_assignments": "One person's share of a task, with their own status and "
                        "progress. A task is not somebody's responsibility until "
                        "they accept it, so this row starts pending.",
    "task_progress_updates": "The progress history of one assignment: from, to, and "
                             "the note that came with it.",
    "task_activities": "Everything else that happened to a task -- assignment, "
                       "reassignment, uploads. Hangs off the task rather than an "
                       "assignment, so it can record an action by someone who is "
                       "not assigned, such as the owner attaching a brief.",
    "task_comments": "Discussion on a task.",
    "task_attachments": "A file submitted as evidence, or handed out with the task. "
                        "The bytes are in the row; see the note below.",
    "work_notifications": "The bell. Work events and personal receipts, with read "
                          "and dismissed state.",
    "push_subscriptions": "One browser or device's Web Push endpoint and keys.",
    "ai_conversations": "Stored assistant transcript.",
}

# Where a column deserves a sentence of its own.
NOTES: dict[tuple[str, str], str] = {
    ("task_attachments", "data"):
        "The file itself. Deferred: never loaded by a listing or dashboard query.",
    ("task_attachments", "size_bytes"): "Capped at 10 MB by the upload endpoint.",
    ("task_assignments", "progress"): "0-100. Reaching 100 sets status to completed.",
    ("task_assignments", "status"):
        "pending / accepted / in_progress / completed / declined.",
    ("work_tasks", "status"): "Derived from the assignments; not set directly.",
    ("departments", "kind"): "'academic' for a teaching department, 'office' for a post.",
    ("users", "password_hash"): "bcrypt. Null for Google-only accounts.",
    ("users", "calendar_token"):
        "Secret path segment for the read-only ICS feed. Rotatable from the profile.",
    ("users", "google_refresh_token"): "Encrypted at rest.",
    ("users", "semester_start"): "Anchors the schedule period on generated documents.",
    ("events", "recurrence_group_id"):
        "Shared by every occurrence of one recurring lecture.",
    ("work_notifications", "kind"):
        "Event type; also what the per-category preference switches filter on.",
}

DERIVED_NOTE = (
    "Two states an owner asks about are not stored anywhere: overdue and "
    "incomplete. They are computed when read, from the assignment's status and "
    "the task's due date. A stored flag would be wrong the moment a clock passed "
    "a due date with nobody looking, and would need a sweep to keep honest."
)


def type_of(column) -> str:
    try:
        return str(column.type)
    except Exception:                                    # dialect-specific types
        return column.type.__class__.__name__


def default_of(column) -> str:
    if column.primary_key:
        return "generated"
    d = column.default
    if d is None:
        return "—"
    if getattr(d, "is_callable", False):
        return "generated"
    arg = getattr(d, "arg", None)
    if arg is None:
        return "—"
    if isinstance(arg, bool):
        return "true" if arg else "false"
    return str(arg)[:40]


def heading(doc, text, size=14, before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p


def body(doc, text, size=9.5, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p


def bullet(doc, text, size=9):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.clear()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def table_of(doc, headers, rows, widths):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    _table_borders(t)
    for i, w in enumerate(widths):
        t.columns[i].width = Inches(w)
    for i, h in enumerate(headers):
        _write_lines(t.rows[0].cells[i], [h], bold=True, size=8.5)
        _shade(t.rows[0].cells[i], HEAD_SHADE)
    for r, row in enumerate(rows, start=1):
        for i, value in enumerate(row):
            _write_lines(t.rows[r].cells[i], [str(value)], size=8.5)
    return t


def document_table(doc, name, table):
    heading(doc, name, size=12, before=13)
    if name in PURPOSE:
        body(doc, PURPOSE[name], size=9)

    rows = []
    for c in table.columns:
        fk = next(iter(c.foreign_keys), None)
        note_parts = []
        if c.primary_key:
            note_parts.append("primary key")
        if fk is not None:
            note_parts.append(f"→ {fk.target_fullname}")
        if c.unique:
            note_parts.append("unique")
        if c.index:
            note_parts.append("indexed")
        extra = NOTES.get((name, c.name))
        if extra:
            note_parts.append(extra)
        rows.append([
            c.name,
            type_of(c),
            "no" if not c.nullable else "yes",
            default_of(c),
            "; ".join(note_parts) or "—",
        ])

    table_of(doc, ["Column", "Type", "Null", "Default", "Notes"],
             rows, [1.5, 1.25, 0.45, 0.9, 2.9])

    # Composite indexes and constraints, which the column table cannot show.
    extras = []
    for idx in sorted(table.indexes, key=lambda i: i.name or ""):
        cols = ", ".join(c.name for c in idx.columns)
        if len(idx.columns) > 1 or idx.unique:
            extras.append(f"{'Unique index' if idx.unique else 'Index'} {idx.name}: ({cols})")
    for con in table.constraints:
        cols = getattr(con, "columns", None)
        if con.__class__.__name__ == "UniqueConstraint" and cols is not None:
            extras.append(f"Unique {con.name}: ({', '.join(c.name for c in cols)})")
    if extras:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        for i, line in enumerate(extras):
            run = p.add_run(("\n" if i else "") + line)
            run.font.size = Pt(8)
            run.italic = True


def build(path: str) -> str:
    md = Base.metadata
    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.7)
    section.top_margin = section.bottom_margin = Inches(0.6)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ProfSchedule AI")
    run.bold = True
    run.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Database schema and logging reference")
    sub_run.font.size = Pt(12)

    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_columns = sum(len(t.columns) for t in md.tables.values())
    stamp_run = stamp.add_run(
        f"Generated {date.today().strftime('%d %B %Y')} from the live models · "
        f"{len(md.tables)} tables · {total_columns} columns"
    )
    stamp_run.italic = True
    stamp_run.font.size = Pt(8.5)

    # ---------------- How the schema is managed ----------------
    heading(doc, "How the schema is managed")
    body(doc,
         "SQLite in development, PostgreSQL in production; the same models drive "
         "both. Tables are created by SQLAlchemy's create_all on startup. There is "
         "no Alembic: existing databases are brought forward by an additive "
         "migration pass that adds any missing column in place, which is why every "
         "new column must be nullable or carry a default.")
    body(doc,
         "One trap is worth recording because it has already caused an outage. "
         "SQLite has no boolean type, so booleans are declared DEFAULT 0 / 1 and "
         "rewritten to FALSE / TRUE for Postgres. That rewrite once ran across "
         "every column definition by plain string replace, which turned INTEGER "
         "DEFAULT 0 into DEFAULT FALSE and INTEGER DEFAULT 120 into DEFAULT "
         "TRUE20 — not valid SQL. The migration threw, rolled every column back "
         "with it, and left production without fields the models selected on every "
         "request. The rewrite now applies only to columns declared BOOLEAN, and a "
         "test walks every registered column to keep it that way.")
    body(doc,
         f"{sum(len(v) for v in _ADDITIVE_COLUMNS.values())} columns across "
         f"{len(_ADDITIVE_COLUMNS)} tables are registered for that pass. A column "
         "added to a model without an entry there works on a fresh database and is "
         "missing on every existing one — which is every real deployment.")

    # ---------------- Replication ----------------
    heading(doc, "The mirror database")
    body(doc,
         "When MIRROR_DATABASE_URL is set, every committed change is copied to a "
         "second database within about a second, and the application switches to "
         "it if the first stops answering. Two infrastructure tables carry that, "
         "in both databases, because either can become the source:")
    bullet(doc,
           "replication_log — one row per changed row: the table, its primary key "
           "and whether it was written or deleted. Appended inside the same "
           "transaction as the change, which is what makes it atomic with the "
           "data and durable across a restart. It records identity, not values, "
           "so a 10 MB attachment does not also become a 13 MB log row.")
    bullet(doc,
           "replication_state — a single row holding the last sequence this "
           "database has applied. The bookmark a restart resumes from.")
    body(doc,
         "This is asynchronous replication: a failover can lose the last second "
         "of writes. Synchronous, zero-loss failover is a database-level feature "
         "— Postgres streaming replication with a coordinator, or a managed "
         "high-availability plan — and no application-level design can honestly "
         "claim it. It also assumes a single writing process; two instances that "
         "disagreed about which database is live would diverge, which is what "
         "REPLICATION_ALLOW_FAILOVER exists to prevent.")

    # ---------------- The tables ----------------
    heading(doc, "Tables")
    documented = set()
    for group_name, blurb, names in GROUPS:
        heading(doc, group_name, size=13, before=16)
        body(doc, blurb, size=9, italic=True)
        for name in names:
            if name in md.tables:
                document_table(doc, name, md.tables[name])
                documented.add(name)

    leftovers = sorted(set(md.tables) - documented)
    if leftovers:
        heading(doc, "Other tables", size=13)
        body(doc, "Present in the models but not yet grouped above.", size=9, italic=True)
        for name in leftovers:
            document_table(doc, name, md.tables[name])

    # ---------------- Logs ----------------
    doc.add_page_break()
    heading(doc, "Logs", size=15, before=0)
    body(doc,
         "The application keeps two different kinds of log, and they answer "
         "different questions. The audit trail lives in the database and is part of "
         "the product: a professor reads it. The runtime log goes to the host's log "
         "stream and is for whoever is keeping the service alive.")

    heading(doc, "Audit trail (in the database)", size=13)
    body(doc,
         "A task's history is assembled at read time from three tables. They are "
         "kept separate because each records a genuinely different thing, and "
         "merging them into one table would mean a row whose meaning depends on "
         "which columns happen to be filled in.")
    for name, what in (
        ("task_progress_updates",
         "Progress on one assignment: from, to, and any note. Written whenever "
         "somebody moves their own slider."),
        ("task_comments", "Discussion, written by any community member."),
        ("task_activities",
         "Everything else: assignment, reassignment, file uploaded, file removed. "
         "Hangs off the task rather than an assignment, which is what lets it "
         "record an action by somebody who is not assigned — an owner attaching a "
         "brief, for instance."),
        ("work_notifications",
         "Delivered notifications and personal receipts, with read and dismissed "
         "state. Also the bell's backing store."),
    ):
        bullet(doc, f"{name} — {what}")
    body(doc,
         "The task detail endpoint merges the first three into a single "
         "newest-first timeline, so the history reads as one story rather than "
         "three tabs the reader has to interleave. Nothing in the audit trail is "
         "ever updated or deleted by the application; rows disappear only when "
         "their task or community is deleted, by cascade.")
    body(doc, DERIVED_NOTE)

    heading(doc, "Runtime logging", size=13)
    body(doc,
         "Configured once in app/main.py at INFO level, formatted as "
         "'%(asctime)s %(levelname)s [%(name)s] %(message)s'. The level is set "
         "explicitly because application loggers otherwise inherit the root's "
         "WARNING and the INFO diagnostics — AI provider fallbacks, reminder "
         "delivery counts — never reach the host's log stream, which on a deployed "
         "instance is the only way to see them.")
    rows = [
        ["app.perf", "WARNING",
         "Requests over 400ms, single queries over 120ms, and any request issuing "
         "more than 25 queries (which usually means one query per row)."],
        ["app.database", "WARNING / ERROR",
         "Configuration problems, index creation failures, and the fallback to "
         "local SQLite when DATABASE_URL is unusable."],
        ["app.main", "ERROR",
         "Startup failures. Also surfaced by /api/health, so a broken deploy "
         "reports why rather than showing an opaque failure."],
        ["app.services.ai_service", "WARNING",
         "Provider errors and falls back to the rule-based parser."],
        ["app.services.work_notify", "INFO / EXCEPTION",
         "Push delivery, and any failure to deliver."],
        ["app.services.reminder_service", "INFO / EXCEPTION",
         "Reminder sweeps and delivery."],
        ["app.routers.cron", "WARNING / ERROR",
         "Refuses to run unauthenticated when CRON_SECRET is unset."],
        ["app.services.google", "WARNING",
         "OAuth exchange and calendar sync failures."],
    ]
    table_of(doc, ["Logger", "Level", "What it records"], rows, [1.9, 1.15, 3.95])
    body(doc,
         "No secret, password, token or file content is ever written to the "
         "runtime log. Slow-query lines truncate the statement at 140 characters "
         "and carry no bound parameters.")

    heading(doc, "Where the bytes live", size=13)
    body(doc,
         "Task attachments are stored in the database rather than on disk or in "
         "object storage. There is no object storage configured for this project, "
         "and the host's filesystem is ephemeral — a file written to it disappears "
         "on the next deploy, silently, which for evidence of somebody's work is "
         "the worst way to lose it. The column is deferred and uploads are capped "
         "at 10 MB, so no listing or dashboard query ever loads a file to show a "
         "filename. Moving to object storage later is this one column and two "
         "endpoints.")

    doc.save(path)
    return path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "ProfSchedule AI - Database and Logs.docx"
    print("written:", build(out))
