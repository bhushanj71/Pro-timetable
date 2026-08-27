"""The weekly timetable as a Word document, in the college's own layout.

The grid is fixed, not derived from the events: ten columns with a short break
after the second period and lunch after the fifth, Monday to Saturday. That is
the form the college hands out and expects back, so the document matches it
whatever shape the professor's own week happens to be.

A fixed grid has one consequence worth being honest about. Anything that does
not fall inside one of these periods -- an evening lab, a Sunday viva -- has no
cell to go in. Rather than dropping it silently, the document lists it under
the table.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]

HEADER_SHADE = "D9D9D9"


@dataclass(frozen=True)
class Slot:
    """One column.

    `label` is reproduced exactly as the college writes it, including the
    "10: 15" and the "10.30". They are inconsistent, and they are what the
    printed form says; a document meant to sit beside the official one should
    not quietly correct it.
    """

    label: str
    start: int | None = None   # minutes from midnight; None on a break column
    end: int | None = None
    break_text: str | None = None


SLOTS: list[Slot] = [
    Slot("8:15 AM to 9:15 AM", 8 * 60 + 15, 9 * 60 + 15),
    Slot("9:15 AM to 10: 15 AM", 9 * 60 + 15, 10 * 60 + 15),
    Slot("10:15 AM to 10:30 AM", break_text="SHORT BREAK"),
    Slot("10.30 AM to 11:30 AM", 10 * 60 + 30, 11 * 60 + 30),
    Slot("11:30 AM to 12:30 PM", 11 * 60 + 30, 12 * 60 + 30),
    Slot("12:30 PM to 01:30 PM", break_text="LUNCH BREAK"),
    Slot("01:30 PM to 02:30 PM", 13 * 60 + 30, 14 * 60 + 30),
    Slot("02:30 PM to 03:30 PM", 14 * 60 + 30, 15 * 60 + 30),
    Slot("03:30 PM to 04:30 PM", 15 * 60 + 30, 16 * 60 + 30),
    Slot("04:30 PM to 05:30 PM", 16 * 60 + 30, 17 * 60 + 30),
]


# --------------------------------------------------------------------------
# Low-level docx helpers.
#
# python-docx exposes no API for cell shading, table borders or vertical text,
# so each of these writes the OOXML element itself. They are small, and they
# are the only place raw XML appears.
# --------------------------------------------------------------------------
def _shade(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _vertical_text(cell) -> None:
    """Bottom-to-top, the way the break columns read on the printed form."""
    direction = OxmlElement("w:textDirection")
    direction.set(qn("w:val"), "btLr")
    cell._tc.get_or_add_tcPr().append(direction)


def _table_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")          # eighths of a point
        el.set(qn("w:color"), "000000")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _write_lines(cell, lines: list[str], *, bold: bool = False, size: int = 9) -> None:
    """Several lines in one cell, as line breaks inside a single paragraph.

    Assigning cell.text first and adding paragraphs after leaves an empty one
    at the top, which shows up as a blank line in every filled cell.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # cell.text = "" leaves an empty run behind, which would sit in front of
    # the first line carrying none of the formatting below.
    for stale in list(para.runs):
        stale._r.getparent().remove(stale._r)

    for i, line in enumerate(lines):
        if i:
            # The break gets a run of its own. Assigning .text replaces
            # everything inside a run, so a break added to the same run before
            # the text is silently wiped by the assignment that follows -- the
            # two lines then render as one.
            para.add_run().add_break()
        run = para.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)


def _columns_for(minutes_start: int, minutes_end: int) -> list[int]:
    """Every teaching column an event overlaps.

    A two-hour lecture occupies two cells rather than being filed under
    whichever one it happens to start in.
    """
    return [
        i for i, slot in enumerate(SLOTS)
        if slot.start is not None
        and minutes_start < slot.end
        and minutes_end > slot.start
    ]


def build_timetable_doc(user, events) -> io.BytesIO:
    doc = Document()

    # Landscape with narrow margins: eleven columns do not fit otherwise.
    section = doc.sections[0]
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.4)
    section.top_margin = section.bottom_margin = Inches(0.5)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(user.name)
    run.bold = True
    run.font.size = Pt(13)

    # Department and college when the work profile is filled in. A timetable
    # handed to an office is identified by more than a name.
    bits = [
        b for b in (
            getattr(getattr(user, "department_rel", None), "name", None),
            getattr(getattr(user, "college_rel", None), "name", None),
        ) if b
    ]
    if bits:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(" · ".join(bits))
        sub_run.font.size = Pt(9)

    table = doc.add_table(rows=len(DAYS) + 1, cols=len(SLOTS) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table)

    # --- Header row ---
    header = table.rows[0]
    _write_lines(header.cells[0], ["Day \\ Time"], bold=True)
    _shade(header.cells[0], HEADER_SHADE)
    for i, slot in enumerate(SLOTS, start=1):
        _write_lines(header.cells[i], [slot.label], bold=True, size=8)
        _shade(header.cells[i], HEADER_SHADE)

    # --- Day column ---
    for r, day in enumerate(DAYS, start=1):
        _write_lines(table.rows[r].cells[0], [day], bold=True)
        _shade(table.rows[r].cells[0], HEADER_SHADE)

    # --- Break columns: one cell spanning every day, text on its side ---
    for i, slot in enumerate(SLOTS, start=1):
        if not slot.break_text:
            continue
        merged = table.cell(1, i).merge(table.cell(len(DAYS), i))
        _write_lines(merged, [slot.break_text], bold=True, size=11)
        _shade(merged, HEADER_SHADE)
        _vertical_text(merged)

    # --- Lectures ---
    grid: dict[tuple[int, int], list[str]] = {}
    unplaced: list[str] = []
    for e in events:
        day = e.start_datetime.weekday()          # Monday is 0, Sunday is 6
        start = e.start_datetime.hour * 60 + e.start_datetime.minute
        end = e.end_datetime.hour * 60 + e.end_datetime.minute
        columns = _columns_for(start, end)

        if day >= len(DAYS) or not columns:
            unplaced.append(
                e.start_datetime.strftime("%A %d %b, %I:%M %p")
                + "–"
                + e.end_datetime.strftime("%I:%M %p")
                + " — "
                + (e.subject or e.title)
            )
            continue

        lines = [e.subject or e.title]
        if e.location:
            lines.append("(" + e.location + ")")
        for c in columns:
            grid.setdefault((day, c), []).extend(lines)

    for (day, col), lines in grid.items():
        _write_lines(table.cell(day + 1, col + 1), lines, size=8)

    # --- Anything the fixed grid has no cell for ---
    if unplaced:
        doc.add_paragraph()
        note = doc.add_paragraph()
        note_run = note.add_run("Outside the timetable hours")
        note_run.bold = True
        note_run.font.size = Pt(9)
        for line in unplaced:
            p = doc.add_paragraph(line, style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
