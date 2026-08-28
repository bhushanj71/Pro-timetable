"""The personalised plan as a Word document.

Deliberately plain. This is a page somebody prints and keeps on a desk, so it
is one column of readable tables with no colour beyond a grey band to separate
what is fixed from what can move. Anything more decorative survives the screen
and dies in a laser printer.

The one visual distinction that earns its place is fixed versus flexible: a
lecture and a study block are not the same kind of thing, and a plan that
renders them identically invites you to move the lecture.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.services.day_planner import DayPlan, span
from app.services.doc_export import _shade, _table_borders, _write_lines

FIXED_SHADE = "E8E8E8"
HEAD_SHADE = "D9D9D9"

# Only where the detail column would otherwise be empty AND the activity name
# does not already say it. "Exercise / Exercise" across two columns is noise,
# but a lecture with no room recorded still wants saying what it is.
KIND_LABEL = {"fixed": "Class / commitment"}


def _hours(minutes: int) -> str:
    if minutes <= 0:
        return "0m"
    h, m = divmod(minutes, 60)
    if not h:
        return f"{m}m"
    return f"{h}h" if not m else f"{h}h {m:02d}m"


def _summary_line(plan: DayPlan) -> str:
    """One sentence a reader can take in without studying the table above it."""
    totals = plan.totals()
    parts = [
        f"{sum(1 for b in plan.blocks if b.kind == 'fixed')} fixed commitment"
        f"{'' if sum(1 for b in plan.blocks if b.kind == 'fixed') == 1 else 's'}",
        f"study {_hours(totals.get('study', 0))}",
        f"breaks {_hours(totals.get('break', 0))}",
    ]
    if totals.get("exercise"):
        parts.append(f"exercise {_hours(totals['exercise'])}")
    if totals.get("travel"):
        parts.append(f"travel {_hours(totals['travel'])}")
    if totals.get("meal"):
        parts.append(f"meals {_hours(totals['meal'])}")

    # Unplanned time inside the day's own span, which is the number people
    # actually want: what is still mine.
    if plan.blocks:
        day_span = max(b.end for b in plan.blocks) - min(b.start for b in plan.blocks)
        free = day_span - sum(b.minutes for b in plan.blocks)
        if free > 0:
            parts.append(f"unplanned {_hours(free)}")
    return " · ".join(parts)


def build_personal_schedule_doc(user, plans: list[DayPlan], period_label: str) -> io.BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.7)
    section.top_margin = section.bottom_margin = Inches(0.6)

    # Body text once, at the top, so every table inherits it.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PERSONALISED SCHEDULE")
    run.bold = True
    run.font.size = Pt(15)

    for label, value in (
        ("Name", user.name),
        ("Schedule period", period_label),
        (
            "Department",
            " · ".join(
                b for b in (
                    getattr(getattr(user, "department_rel", None), "name", None),
                    getattr(getattr(user, "college_rel", None), "name", None),
                ) if b
            ),
        ),
    ):
        if not value:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        head = p.add_run(f"{label}: ")
        head.bold = True
        head.font.size = Pt(10)
        p.add_run(str(value)).font.size = Pt(10)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note_run = note.add_run(
        "Shaded rows are fixed commitments from your timetable. Everything else "
        "is planned around them and can move."
    )
    note_run.italic = True
    note_run.font.size = Pt(8.5)

    if not plans:
        empty = doc.add_paragraph()
        empty.add_run(
            "There is nothing in your timetable for this period, so there is "
            "nothing to plan around yet. Add your classes and generate this again."
        ).font.size = Pt(10)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    for plan in plans:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        h = heading.add_run(f"{plan.name} — {plan.day.strftime('%d %b %Y')}")
        h.bold = True
        h.font.size = Pt(12)

        table = doc.add_table(rows=len(plan.blocks) + 1, cols=3)
        _table_borders(table)
        table.columns[0].width = Inches(1.6)
        table.columns[1].width = Inches(3.2)
        table.columns[2].width = Inches(2.3)

        for i, label in enumerate(("Time", "Activity", "Location / details")):
            _write_lines(table.rows[0].cells[i], [label], bold=True, size=9)
            _shade(table.rows[0].cells[i], HEAD_SHADE)

        for r, block in enumerate(plan.blocks, start=1):
            cells = table.rows[r].cells
            _write_lines(cells[0], [span(block.start, block.end)], size=9)
            _write_lines(cells[1], [block.title], bold=block.kind == "fixed", size=9)
            detail = block.detail or KIND_LABEL.get(block.kind, "")
            _write_lines(cells[2], [detail], size=9)
            if block.kind == "fixed":
                for c in cells:
                    _shade(c, FIXED_SHADE)

        summary = doc.add_paragraph()
        summary.paragraph_format.space_before = Pt(3)
        s = summary.add_run(_summary_line(plan))
        s.font.size = Pt(8.5)
        s.italic = True

        # Why the day looks the way it does. Without this a moved lunch reads
        # as the planner being wrong rather than the day being full.
        for line in plan.notes:
            n = doc.add_paragraph()
            n.paragraph_format.space_after = Pt(0)
            run = n.add_run(f"— {line}")
            run.font.size = Pt(8.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
