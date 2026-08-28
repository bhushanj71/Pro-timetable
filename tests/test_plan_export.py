"""The personalised schedule document."""
import io

from docx import Document


def _prefs(client, **over):
    body = {
        "day_start": "07:00", "day_end": "22:30",
        "lunch_start": "13:00", "lunch_end": "13:30",
        "study_target_minutes": 120, "exercise_minutes": 30,
        "subject_priorities": "ANN, DBMS",
    }
    body.update(over)
    r = client.put("/api/auth/me", json=body)
    assert r.status_code == 200, r.text
    return r.json()


def test_the_plan_downloads_as_a_word_document(auth_client):
    _prefs(auth_client)
    r = auth_client.get("/api/export/plan")
    assert r.status_code == 200
    assert r.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert r.content[:2] == b"PK"
    doc = Document(io.BytesIO(r.content))
    assert "PERSONALISED SCHEDULE" in "\n".join(p.text for p in doc.paragraphs)


def test_the_document_names_the_professor_and_the_period(auth_client):
    _prefs(auth_client)
    doc = Document(io.BytesIO(auth_client.get("/api/export/plan").content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Dr. Test" in text
    assert "Schedule period" in text


def test_a_semester_start_dates_the_term_and_numbers_the_week(auth_client):
    _prefs(auth_client, semester_start="2026-01-05")
    doc = Document(io.BytesIO(auth_client.get("/api/export/plan").content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "week" in text and "of term" in text


def test_an_empty_timetable_says_so_instead_of_inventing_a_week(auth_client):
    """The spec's own rule: never produce a generic schedule. With nothing
    uploaded there is nothing to plan around, and saying so is the honest
    output -- a plausible invented week would be worse than none."""
    _prefs(auth_client, working_days="")
    doc = Document(io.BytesIO(auth_client.get("/api/export/plan").content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "nothing to plan around yet" in text
    assert not doc.tables


def test_the_days_carry_a_three_column_table(auth_client):
    _prefs(auth_client)
    auth_client.post("/api/events", json={
        "title": "ANN Lecture", "subject": "ANN", "event_type": "lecture",
        "start_datetime": "2026-09-07T04:30:00Z", "end_datetime": "2026-09-07T05:30:00Z",
    })
    doc = Document(io.BytesIO(auth_client.get("/api/export/plan").content))
    assert doc.tables, "a working week should produce at least one day table"
    header = [c.text for c in doc.tables[0].rows[0].cells]
    assert header == ["Time", "Activity", "Location / details"]


def test_the_semester_date_is_saved_from_the_profile_endpoint(auth_client):
    r = auth_client.put("/api/auth/me", json={"semester_start": "2026-07-15"})
    assert r.status_code == 200
    assert r.json()["semester_start"] == "2026-07-15"
    assert auth_client.get("/api/auth/me").json()["semester_start"] == "2026-07-15"


def test_nonsense_planning_values_are_refused_at_the_edge(auth_client):
    """Bounded in the schema so the planner never has to defend itself against
    a nine-hour study block or a negative commute."""
    assert auth_client.put("/api/auth/me", json={"study_block_min": 5}).status_code == 422
    assert auth_client.put("/api/auth/me", json={"commute_minutes": -10}).status_code == 422
    assert auth_client.put("/api/auth/me", json={"focus_period": "whenever"}).status_code == 422
    # min longer than max is individually valid and jointly nonsense.
    assert auth_client.put(
        "/api/auth/me", json={"study_block_min": 200, "study_block_max": 60}
    ).status_code == 422


def test_the_timetable_page_offers_the_download_and_the_semester_date(auth_client):
    page = auth_client.get("/timetable").text
    assert 'id="export-plan-btn"' in page
    assert "Download Personalized Schedule" in page
    assert 'id="tt-semester-start"' in page
