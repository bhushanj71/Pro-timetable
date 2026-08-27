"""The timetable as the college's own Word form.

The grid is fixed and the college's, not ours, so most of these assert the
shape of the document rather than anything about the events: the labels
verbatim, the two breaks merged down the page with their text on its side, and
Monday to Saturday. If any of that drifts, the document stops being the form
the office accepts, which is the only thing it is for.
"""
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn

from app.services.doc_export import DAYS, SLOTS, build_timetable_doc


class _User:
    name = "Dr. Test"
    department_rel = None
    college_rel = None


class _Event:
    def __init__(self, day, start, end, subject, location=None):
        # 2026-08-24 is a Monday, so day 0 lands on Monday.
        self.start_datetime = datetime(2026, 8, 24 + day, *start)
        self.end_datetime = datetime(2026, 8, 24 + day, *end)
        self.subject = subject
        self.title = subject
        self.location = location


def _doc(events):
    return Document(build_timetable_doc(_User(), events))


def _table(events):
    return _doc(events).tables[0]


# --------------------------------------------------------------------------
# The form itself
# --------------------------------------------------------------------------
def test_the_grid_is_the_colleges_grid():
    t = _table([])
    assert len(t.rows) == len(DAYS) + 1
    assert len(t.columns) == len(SLOTS) + 1


def test_the_period_labels_are_reproduced_verbatim():
    """Including "10: 15" and "10.30". They are inconsistent, and they are what
    the printed form says; correcting them would make this a different form."""
    header = [c.text for c in _table([]).rows[0].cells]
    assert header == [
        "Day \\ Time",
        "8:15 AM to 9:15 AM",
        "9:15 AM to 10: 15 AM",
        "10:15 AM to 10:30 AM",
        "10.30 AM to 11:30 AM",
        "11:30 AM to 12:30 PM",
        "12:30 PM to 01:30 PM",
        "01:30 PM to 02:30 PM",
        "02:30 PM to 03:30 PM",
        "03:30 PM to 04:30 PM",
        "04:30 PM to 05:30 PM",
    ]


def test_the_week_runs_monday_to_saturday():
    t = _table([])
    assert [r.cells[0].text for r in t.rows[1:]] == [
        "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday",
    ]


def test_the_breaks_span_every_day_with_their_text_on_its_side():
    t = _table([])
    for column, label in ((3, "SHORT BREAK"), (6, "LUNCH BREAK")):
        cell = t.cell(1, column)
        assert cell.text == label
        # One cell, reached from all six day rows -- that is what merged means.
        spanned = sum(1 for r in t.rows[1:] if r.cells[column]._tc is cell._tc)
        assert spanned == len(DAYS), f"{label} spans {spanned} rows, not {len(DAYS)}"

        direction = cell._tc.tcPr.find(qn("w:textDirection"))
        assert direction is not None and direction.get(qn("w:val")) == "btLr"


def test_the_headers_and_breaks_are_shaded():
    t = _table([])
    shaded = [t.cell(0, 0), t.cell(0, 1), t.cell(1, 0), t.cell(1, 3), t.cell(1, 6)]
    for cell in shaded:
        shd = cell._tc.tcPr.find(qn("w:shd"))
        assert shd is not None and shd.get(qn("w:fill")) == "D9D9D9"
    # A period cell is not.
    assert t.cell(1, 1)._tc.tcPr is None or \
        t.cell(1, 1)._tc.tcPr.find(qn("w:shd")) is None


def test_every_cell_is_ruled():
    borders = _table([])._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    edges = {e.tag.split("}")[1] for e in borders}
    assert edges == {"top", "left", "bottom", "right", "insideH", "insideV"}


def test_the_page_is_landscape():
    """Eleven columns do not fit on a portrait page."""
    section = _doc([]).sections[0]
    assert section.page_width > section.page_height


# --------------------------------------------------------------------------
# Where the lectures land
# --------------------------------------------------------------------------
def test_a_lecture_lands_in_its_own_period():
    t = _table([_Event(0, (8, 15), (9, 15), "ANN")])
    assert "ANN" in t.cell(1, 1).text
    assert t.cell(1, 2).text == ""


def test_an_afternoon_lecture_clears_the_lunch_column():
    t = _table([_Event(0, (13, 30), (14, 30), "DSA")])
    assert "DSA" in t.cell(1, 7).text
    assert t.cell(1, 6).text == "LUNCH BREAK"


def test_a_two_hour_lab_occupies_both_periods():
    """Filing it under whichever period it starts in would hide half of it."""
    t = _table([_Event(3, (14, 30), (16, 30), "AI Lab")])
    assert "AI Lab" in t.cell(4, 8).text
    assert "AI Lab" in t.cell(4, 9).text


def test_the_room_is_on_its_own_line():
    """A run's .text setter replaces everything in the run, so a line break
    added to the same run is wiped by it -- and the room ends up jammed onto
    the end of the subject."""
    cell = _table([_Event(0, (8, 15), (9, 15), "ANN", "Lab 402")]).cell(1, 1)
    assert "w:br" in cell._tc.xml
    assert [r.text for p in cell.paragraphs for r in p.runs] == ["ANN", "\n", "(Lab 402)"]


def test_two_lectures_in_one_period_both_appear():
    cell = _table([
        _Event(0, (8, 15), (9, 15), "ANN"),
        _Event(0, (8, 15), (9, 15), "Tutorial"),
    ]).cell(1, 1)
    assert "ANN" in cell.text and "Tutorial" in cell.text


# --------------------------------------------------------------------------
# What the fixed grid has no room for
# --------------------------------------------------------------------------
def test_an_evening_class_is_listed_rather_than_dropped():
    """The grid stops at 5:30pm. Silently losing a lecture from a document
    someone hands to an office is worse than the grid being fixed."""
    doc = _doc([_Event(4, (19, 0), (20, 0), "Evening Doubt Session")])
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Outside the timetable hours" in text
    assert "Evening Doubt Session" in text


def test_a_sunday_class_is_listed_rather_than_dropped():
    doc = _doc([_Event(6, (11, 0), (12, 0), "Sunday Viva")])
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sunday Viva" in text


def test_nothing_is_said_about_hours_when_everything_fits():
    doc = _doc([_Event(0, (8, 15), (9, 15), "ANN")])
    assert "Outside the timetable hours" not in "\n".join(p.text for p in doc.paragraphs)


# --------------------------------------------------------------------------
# The route, and the page that offers it
# --------------------------------------------------------------------------
def test_the_endpoint_returns_a_word_document(auth_client):
    r = auth_client.get("/api/export/doc")
    assert r.status_code == 200
    assert r.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert ".docx" in r.headers["content-disposition"]
    # A real docx, not an error page with the right header on it.
    assert r.content[:2] == b"PK"
    assert len(Document(__import__("io").BytesIO(r.content)).tables) == 1


def test_the_timetable_page_offers_one_export(auth_client):
    page = auth_client.get("/timetable").text
    assert 'id="export-doc-btn"' in page
    for gone in ("export-csv-btn", "export-ics-btn", "export-pdf-btn", "open-generator-btn"):
        assert f'id="{gone}"' not in page, gone


def test_the_other_export_routes_still_work(auth_client):
    """They are off the timetable menu, not deleted -- the calendar
    subscription feed is built on the ICS one."""
    for path in ("/api/export/csv", "/api/export/ics"):
        assert auth_client.get(path).status_code == 200
