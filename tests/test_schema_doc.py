"""The schema document must not drift from the schema.

The document is generated from the live metadata, so its column lists cannot
lie. What it can do is quietly file a new table under "Other tables" at the
back, where nobody looks. This keeps every table deliberately placed.
"""
import app.models  # noqa: F401  -- registers the tables
from app.database import Base
from tools.schema_doc import GROUPS, PURPOSE


GROUPED = {name for _, _, names in GROUPS for name in names}


def test_every_table_is_placed_in_a_group():
    """A new table should be filed on purpose, not swept into the appendix."""
    missing = set(Base.metadata.tables) - GROUPED
    assert not missing, (
        f"add {sorted(missing)} to a group in tools/schema_doc.py "
        "so it appears where a reader would look for it"
    )


def test_no_group_names_a_table_that_does_not_exist():
    """A renamed or dropped table would otherwise leave a heading behind with
    nothing under it."""
    stale = GROUPED - set(Base.metadata.tables)
    assert not stale, f"tools/schema_doc.py still lists {sorted(stale)}"


def test_every_table_has_a_stated_purpose():
    """Introspection can list a column. Only the prose says why the table is
    there, and a table with no purpose written down is one nobody can review."""
    missing = set(Base.metadata.tables) - set(PURPOSE)
    assert not missing, f"describe {sorted(missing)} in PURPOSE"


def test_the_document_builds_and_covers_everything(tmp_path):
    from docx import Document

    from tools.schema_doc import build

    out = build(str(tmp_path / "schema.docx"))
    doc = Document(out)

    headings = {p.text.strip() for p in doc.paragraphs}
    assert set(Base.metadata.tables) <= headings

    documented = {
        row.cells[0].text.strip()
        for table in doc.tables
        for row in table.rows[1:]
    }
    for name, table in Base.metadata.tables.items():
        for column in table.columns:
            assert column.name in documented, f"{name}.{column.name} is not in the document"
